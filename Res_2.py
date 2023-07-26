import os
import time
import sys
import keyboard
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import smtplib
from email.message import EmailMessage
from Resume import name, email, phone, objective, experience, skills, education, certificates_training, additional_skills
from env import SMTP_SERVER, SMTP_PORT, SMTP_USERNAME, SMTP_PASSWORD

# Function to clear the screen for better user experience
def clear_screen():
    os.system('cls' if os.name == 'nt' else 'clear')

# Function to return to the menu after an operation
def return_to_menu():
    input("Press Enter to return to the menu...")
    clear_screen()
    interactive_menu()

# Function to generate the PDF resume
def generate_pdf_resume():
    pdf_filename = "Resume.pdf"
    page_width, page_height = letter
    padding_top = 25  # Adjusted padding at the top
    padding_bottom = 25  # Adjusted padding at the bottom

    c = canvas.Canvas(pdf_filename, pagesize=(page_width, page_height))

    # Set the font size and leading (line spacing)
    font_size = 12
    leading = 14

    # Set the starting position for the content
    x = 50
    y = page_height - padding_top  # Start from the top with padding

    # Write the resume details to the PDF
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, name)
    y -= leading
    c.setFont("Helvetica", font_size)
    c.drawString(x, y, email)
    y -= leading
    c.drawString(x, y, phone)
    y -= 3 * leading

    # Function to write text and calculate the remaining y position
    def write_text(text, line_spacing=leading):
        nonlocal y
        lines = text.strip().split("\n")
        for line in lines:
            c.drawString(x, y, line)
            y -= line_spacing

    # Write the objective section
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, "Objective:")
    y -= leading
    c.setFont("Helvetica", font_size)
    write_text(objective, leading + 2)  # Increased line spacing after the objective section

    # Write the work experience section
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, "Work Experience:")
    y -= leading
    c.setFont("Helvetica", font_size)
    for exp in experience:
        c.drawString(x, y, f"Title: {exp['title']}")
        y -= leading
        c.drawString(x, y, f"Company: {exp['company']}")
        y -= leading
        c.drawString(x, y, f"Date: {exp['date']}")
        y -= leading
        write_text(exp['description'], leading + 2)  # Increased line spacing after each experience entry

    # Write the skills section
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, "Skills:")
    y -= leading
    c.setFont("Helvetica", font_size)
    for skill, description in skills.items():
        c.drawString(x, y, f"- {skill}: {description}")
        y -= leading

    # Write the education section
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, "Education:")
    y -= leading
    c.setFont("Helvetica", font_size)
    for edu in education:
        c.drawString(x, y, f"School: {edu['school']}")
        y -= leading
        c.drawString(x, y, f"Location: {edu['location']}")
        y -= leading
        c.drawString(x, y, f"Degree: {edu['degree']}")
        y -= leading
        c.drawString(x, y, f"Date: {edu['date']}")
        y -= leading
        write_text(edu['achievements'], leading + 2)  # Increased line spacing after each education entry

    # Write the certificates / training section
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, "Certificates / Training:")
    y -= leading
    c.setFont("Helvetica", font_size)
    for cert in certificates_training:
        c.drawString(x, y, f"- {cert}")
        y -= leading

    # Write the additional skills section
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, "Additional Skills:")
    y -= leading
    c.setFont("Helvetica", font_size)
    for skill in additional_skills:
        c.drawString(x, y, f"- {skill}")
        y -= leading

    # Calculate the required page height for the remaining content
    total_content_height = page_height - y - padding_bottom
    if total_content_height < 0:
        total_content_height = 0

    # Continue writing on a new page if the content does not fit on the current page
    if total_content_height > page_height - padding_bottom:
        c.showPage()
        y = page_height - padding_top

    # Save the PDF
    c.save()
    print(f"PDF resume generated as 'Resume.pdf'")
    clear_screen()
    interactive_menu()

# Function to generate the DOCX resume
def generate_docx_resume():
    doc_filename = "Resume.docx"
    document = Document()

    # Add the resume details to the document
    document.add_heading("Resume Details", level=1)
    document.add_paragraph(f"Name: {name}")
    document.add_paragraph(f"Email: {email}")
    document.add_paragraph(f"Phone: {phone}\n")

    # Add the objective section
    document.add_heading("Objective", level=2)
    document.add_paragraph(objective.strip())

    # Add the work experience section
    document.add_heading("Work Experience", level=2)
    for exp in experience:
        document.add_heading(exp['title'], level=3)
        document.add_paragraph(f"Company: {exp['company']}")
        document.add_paragraph(f"Date: {exp['date']}")
        document.add_paragraph(exp['description'])

    # Add the skills section
    document.add_heading("Skills", level=2)
    for skill, description in skills.items():
        document.add_paragraph(f"- {skill}: {description}")

    # Add the education section
    document.add_heading("Education", level=2)
    for edu in education:
        document.add_heading(edu['school'], level=3)
        document.add_paragraph(f"Location: {edu['location']}")
        document.add_paragraph(f"Degree: {edu['degree']}")
        document.add_paragraph(f"Date: {edu['date']}")
        document.add_paragraph(edu['achievements'])

    # Add the certificates / training section
    document.add_heading("Certificates / Training", level=2)
    for cert in certificates_training:
        document.add_paragraph(f"- {cert}")

    # Add the additional skills section
    document.add_heading("Additional Skills", level=2)
    for skill in additional_skills:
        document.add_paragraph(f"- {skill}")

    # Save the document
    document.save(doc_filename)
    print(f"DOCX resume generated as '{doc_filename}'")
    clear_screen()
    interactive_menu()
    
# Function to generate the TXT resume
def generate_txt_resume():
    txt_filename = "Resume.txt"
    with open(txt_filename, "w") as txt_file:
        # Write the resume details to the file
        txt_file.write("Resume Details:\n")
        txt_file.write(f"Name: {name}\n")
        txt_file.write(f"Email: {email}\n")
        txt_file.write(f"Phone: {phone}\n\n")

        # Write the objective section
        txt_file.write("Objective:\n")
        txt_file.write(objective.strip() + "\n\n")

        # Write the work experience section
        txt_file.write("Work Experience:\n")
        for exp in experience:
            txt_file.write(f"\nTitle: {exp['title']}\n")
            txt_file.write(f"Company: {exp['company']}\n")
            txt_file.write(f"Date: {exp['date']}\n")
            txt_file.write(exp['description'] + "\n")
        txt_file.write("\n")

        # Write the skills section
        txt_file.write("Skills:\n")
        for skill, description in skills.items():
            txt_file.write(f"\n- {skill}: {description}\n")
        txt_file.write("\n")

        # Write the education section
        txt_file.write("Education:\n")
        for edu in education:
            txt_file.write(f"\nSchool: {edu['school']}\n")
            txt_file.write(f"Location: {edu['location']}\n")
            txt_file.write(f"Degree: {edu['degree']}\n")
            txt_file.write(f"Date: {edu['date']}\n")
            txt_file.write(edu['achievements'] + "\n")
        txt_file.write("\n")

        # Write the certificates / training section
        txt_file.write("Certificates / Training:\n")
        for cert in certificates_training:
            txt_file.write(f"\n- {cert}\n")
        txt_file.write("\n")

        # Write the additional skills section
        txt_file.write("Additional Skills:\n")
        for skill in additional_skills:
            txt_file.write(f"\n- {skill}\n")
        txt_file.write("\n")

    print(f"TXT resume generated as '{txt_filename}'")
    clear_screen()
    interactive_menu()
        
# Function to send the resume via email
def send_resume_email(to_email):
    try:
        # Generate the PDF resume
        generate_pdf_resume()

        # Email content
        msg = EmailMessage()
        msg["Subject"] = f"{name} - Resume"
        msg["From"] = SMTP_USERNAME
        msg["To"] = to_email

        email_body = f"""
        Dear Hiring Manager,

        I am excited to submit my resume for your review. Please find attached my resume.

        This resume has been generated by a Python program that kcarden wrote for this purpose. You can view the source code on my GitHub repository, here: 
        
        https://github.com/kcarden/PythonResume

        Thank you for considering my application.

        Best regards,
        {name}
        """

        msg.set_content(email_body)


        # Attach the generated PDF resume to the email
        with open("Resume.pdf", "rb") as file:
            msg.add_attachment(file.read(), maintype="application", subtype="pdf", filename="Resume.pdf")

        # Connect to the SMTP server and send the email
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USERNAME, SMTP_PASSWORD)
            server.send_message(msg)

        print("Resume sent successfully!")

        # Remove the PDF file and return to main menu after sending the email
        print("Resume sent successfully!")
        os.remove("Resume.pdf")
        print("PDF file removed.")
    except Exception as e:
        print(f"Error while sending the resume: {e}")
    clear_screen()
    interactive_menu()
    
# Function to display the entire resume from Resume.py
def display_resume():
    print("\nResume Details:")
    print(f"Name: {name}")
    print(f"Email: {email}")
    print(f"Phone: {phone}\n")

    print("Objective:")
    print(objective.strip() + "\n")

    print("Work Experience:")
    for exp in experience:
        print(f"\nTitle: {exp['title']}")
        print(f"Company: {exp['company']}")
        print(f"Date: {exp['date']}")
        print(exp['description'])
    print()

    print("Skills:")
    for skill, description in skills.items():
        print(f"\n- {skill}: {description}")
    print()

    print("Education:")
    for edu in education:
        print(f"\nSchool: {edu['school']}")
        print(f"Location: {edu['location']}")
        print(f"Degree: {edu['degree']}")
        print(f"Date: {edu['date']}")
        print(edu['achievements'])
    print()

    print("Certificates / Training:")
    for cert in certificates_training:
        print(f"\n- {cert}")
    print()

    print("Additional Skills:")
    for skill in additional_skills:
        print(f"\n- {skill}")
    print()

    # Wait for the user to press Enter to return to the menu
    input("Press Enter to return to the menu...")
    clear_screen()
    interactive_menu()

# Function to exit the program
def exit_program():
    print("Exiting the program...")
    sys.exit(0)

# Interactive menu for options
def interactive_menu():
    while True:
        print("\nInteractive Resume Viewer")
        print("1. Display Resume")
        print("2. Email Resume")
        print("3. Generate PDF Resume")
        print("4. Generate DOCX Resume")
        print("5. Generate TXT Resume")
        print("6. Exit")

        choice = input("Enter your choice (1/2/3/4/5/6): ")

        if choice == "1":
            display_resume()
        elif choice == "2":
            email = input("Enter the recipient's email address: ")
            send_resume_email(email)
        elif choice == "3":
            generate_pdf_resume()
        elif choice == "4":
            generate_docx_resume()
        elif choice == "5":
            generate_txt_resume()
        elif choice == "6":
            exit_program()
        else:
            print("Invalid choice. Please try again.")
            
if __name__ == "__main__":
    working_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(working_dir)

    interactive_menu()